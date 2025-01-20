import os
import json
import streamlit as st
from groq import Groq
import PyPDF2
import docx
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import re
import difflib
import nltk
from nltk.translate.bleu_score import sentence_bleu
from rouge_score import rouge_scorer
from nltk.translate.bleu_score import SmoothingFunction

# Streamlit page configuration
st.set_page_config(
    page_title="LLAMA 3.1. Chat",
    page_icon="🦙",
    layout="centered"
)

@st.cache_data
def download_nltk_resources():
    """
    Comprehensive NLTK resource download
    """
    try:
        # Download specific resources explicitly
        nltk.download('punkt', quiet=True)
        nltk.download('punkt_tab', quiet=True)
        
        # Additional resources you might need
        nltk.download('wordnet', quiet=True)
        nltk.download('stopwords', quiet=True)
        
        return True
    except Exception as e:
        st.error(f"NLTK Resource Download Failed: {e}")
        return False

# Call the download function
download_nltk_resources()

# Alternative tokenization method as a fallback
def safe_tokenize(text):
    """
    Safe tokenization with multiple fallback methods
    """
    try:
        # Try NLTK tokenization first
        return nltk.word_tokenize(text)
    except Exception:
        try:
            # Fallback to simple split
            return text.split()
        except Exception:
            # Last resort: character-level tokenization
            return list(text)


working_dir = os.path.dirname(os.path.abspath(__file__))
config_data = json.load(open(f"{working_dir}/config.json"))

# Save the API key to environment variable
os.environ["GROQ_API_KEY"] = config_data["GROQ_API_KEY"]

# Use the API key diry when initializing the client
client = Groq(api_key=os.environ["GROQ_API_KEY"])

# Initialize the summary history if not present
if "summary_history" not in st.session_state:
    st.session_state.summary_history = []

# Initialize summaries if not present
if "summaries" not in st.session_state:
    st.session_state.summaries = {}

if "uploaded_files_content" not in st.session_state:
    st.session_state.uploaded_files_content = {}

# Streamlit page title
st.title("📚 Multi-Document Research Paper Summarizer")


# Function to extract text from PDF
def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Add these new helper functions
def get_alternative_models():
    """Get a list of alternative models to use when primary model hits rate limit"""
    return {
        "primary": "llama-3.2-11b-vision-preview",
        "backups": [
            "llama-3.2-3b-preview",
            "llama-3.1-70b-versatile",
            "llama-3.1-8b-instant"
        ]
    }

def batch_documents(documents, batch_size=2):
    """Split documents into batches for parallel processing"""
    for i in range(0, len(documents), batch_size):
        yield documents[i:i + batch_size]

def process_document(pdf_content, model, client):
    """Process a single document"""
    try:
        messages = [
            {"role": "system", "content": "Summarize this document focusing on key points."},
            {"role": "user", "content": pdf_content[:4000]}  # First 4000 chars for initial summary
        ]
        
        response = client.chat.completions.create(
            model=model,
            messages=messages
        )
        return response.choices[0].message.content
    except Exception as e:
        st.warning(f"Error with {model}, trying fallback model...")
        return None

def create_cross_document_summary(summaries, original_docs, model, client):
    """
    Improved cross-document summary generation
    """
    try:
        # First, identify key themes across documents
        theme_prompt = """Analyze these document summaries and identify:
        1. Common themes and topics
        2. Shared methodologies
        3. Related findings
        4. Contradictions or disagreements
        5. Knowledge gaps
        
        Format the response as a structured list."""
        
        theme_messages = [
            {"role": "system", "content": theme_prompt},
            {"role": "user", "content": "\n\n".join(summaries)}
        ]
        
        theme_response = client.chat.completions.create(
            model=model,
            messages=theme_messages
        )
        
        themes = theme_response.choices[0].message.content
        
        # Then, generate the cross-document summary using the identified themes
        summary_prompt = f"""Based on the following themes:

        {themes}

        Create a comprehensive cross-document summary that:
        1. Synthesizes the main findings while maintaining accuracy
        2. Highlights relationships between different research aspects
        3. Discusses contradictions and agreements
        4. Presents a coherent narrative of the collective research
        5. Maintains technical accuracy and proper attribution
        
        Focus on creating a summary that would be useful for researchers in this field."""
        
        summary_messages = [
            {"role": "system", "content": summary_prompt},
            {"role": "user", "content": "\n\n".join(summaries)}
        ]
        
        summary_response = client.chat.completions.create(
            model=model,
            messages=summary_messages
        )
        
        return summary_response.choices[0].message.content
        
    except Exception as e:
        st.error(f"Error in cross-document summary: {str(e)}")
        return "\n\n".join(summaries)

def validate_summary(original_documents, generated_summary, individual_summaries):
    """
    Enhanced summary validation with improved comparison logic
    
    Args:
        original_documents (dict): Original documents
        generated_summary (str): Generated cross-document summary
        individual_summaries (dict): Individual document summaries
    
    Returns:
        dict: Validation metrics
    """
    try:
        # Initialize ROUGE scorer with different variants
        rouge_scorer_obj = rouge_scorer.RougeScorer(
            ['rouge1', 'rouge2', 'rougeL'], 
            use_stemmer=True
        )
        
        # Separate validation for individual summaries and cross-document summary
        individual_scores = {
            'rouge1': [], 'rouge2': [], 'rougeL': [], 'bleu': []
        }
        
        cross_doc_scores = {
            'rouge1': [], 'rouge2': [], 'rougeL': [], 'bleu': []
        }
        
        # 1. Validate individual summaries against their source documents
        for doc_name, original_text in original_documents.items():
            if doc_name in individual_summaries:
                individual_summary = individual_summaries[doc_name]
                
                # ROUGE scores for individual summary
                rouge_result = rouge_scorer_obj.score(original_text, individual_summary)
                individual_scores['rouge1'].append(rouge_result['rouge1'].fmeasure)
                individual_scores['rouge2'].append(rouge_result['rouge2'].fmeasure)
                individual_scores['rougeL'].append(rouge_result['rougeL'].fmeasure)
                
                # BLEU score for individual summary
                individual_scores['bleu'].append(
                    calculate_bleu_score(original_text, individual_summary)
                )
        
        # 2. Validate cross-document summary
        # Compare against concatenated individual summaries instead of original docs
        combined_summaries = " ".join(individual_summaries.values())
        
        # ROUGE scores for cross-document summary
        cross_rouge = rouge_scorer_obj.score(combined_summaries, generated_summary)
        cross_doc_scores['rouge1'] = [cross_rouge['rouge1'].fmeasure]
        cross_doc_scores['rouge2'] = [cross_rouge['rouge2'].fmeasure]
        cross_doc_scores['rougeL'] = [cross_rouge['rougeL'].fmeasure]
        
        # BLEU score for cross-document summary
        cross_doc_scores['bleu'] = [
            calculate_bleu_score(combined_summaries, generated_summary)
        ]
        
        # Calculate averages
        avg_individual_scores = {
            metric: sum(scores)/len(scores) if scores else 0
            for metric, scores in individual_scores.items()
        }
        
        avg_cross_doc_scores = {
            metric: sum(scores)/len(scores) if scores else 0
            for metric, scores in cross_doc_scores.items()
        }
        
        return {
            'individual_summaries': {
                'average_scores': avg_individual_scores,
                'detailed_scores': individual_scores
            },
            'cross_document_summary': {
                'average_scores': avg_cross_doc_scores,
                'detailed_scores': cross_doc_scores
            }
        }
        
    except Exception as e:
        st.error(f"Validation error: {str(e)}")
        return None

def calculate_bleu_score(reference_text, candidate_text):
    """
    Improved BLEU score calculation
    """
    try:
        # Improved tokenization
        def tokenize_better(text):
            # Remove special characters and normalize whitespace
            text = re.sub(r'[^\w\s]', ' ', text)
            text = ' '.join(text.split())
            return text.lower().split()
        
        # Prepare reference and candidate
        reference_tokens = tokenize_better(reference_text)
        candidate_tokens = tokenize_better(candidate_text)
        
        # Use multiple reference sentences for better BLEU calculation
        reference_sentences = [reference_tokens[i:i+20] for i in range(0, len(reference_tokens), 10)]
        if not reference_sentences:
            reference_sentences = [reference_tokens]
        
        # Calculate BLEU with smoothing
        weights = (0.25, 0.25, 0.25, 0.25)  # Equal weights for 1-4 grams
        score = sentence_bleu(
            reference_sentences,
            candidate_tokens,
            weights=weights,
            smoothing_function=SmoothingFunction().method1
        )
        
        return score
        
    except Exception as e:
        print(f"BLEU score calculation error: {e}")
        return 0

# Add these new functions for key work analysis

def extract_key_findings(text, model, client):
    """Extract key findings and contributions from the text"""
    try:
        messages = [
            {"role": "system", "content": """You are an expert paralegal with an experience of over 18 years at the top government owned research firm. You are well known for your detailed summarization with to the point. The summary you generate should cover all the things to know in the documents also whenever you refer to any of the provided document it should start with "In the document 'title of the document' the 'Last Name of the First author of the document' et al ". Analyze this research text and extract:
                1. Main research contributions
                2. Key methodologies used
                3. Important findings and results
                4. Novel approaches or techniques introduced
                5. Technical innovations
                Please format the response in clear sections."""},
            {"role": "user", "content": text[:4000]}
        ]
        
        response = client.chat.completions.create(
            model=model,
            messages=messages
        )
        return response.choices[0].message.content
    except Exception as e:
        st.warning(f"Error extracting key findings: {str(e)}")
        return None

def analyze_research_impact(summaries, model, client):
    """Analyze the collective research impact and innovations"""
    try:
        combined_analysis = "\n\n".join(summaries)
        messages = [
            {"role": "system", "content": """You are an expert paralegal with an experience of over 18 years at the top government owned research firm. You are well known for your detailed summarization with to the point. The summary you generate should cover all the things to know in the documents also whenever you refer to any of the provided document it should start with "In the document 'title of the document' the 'Last Name of the First author of the document' et al ". Create a comprehensive analysis of the research contributions that:
                1. Identifies major technical innovations
                2. Highlights novel methodologies
                3. Summarizes key experimental results
                4. Points out unique approaches
                5. Describes practical applications
                6. Notes potential future research directions
                
                Format the response with clear headings for each category."""},
            {"role": "user", "content": combined_analysis}
        ]
        
        response = client.chat.completions.create(
            model=model,
            messages=messages
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error in research impact analysis: {str(e)}")
        return "\n\n".join(summaries)

# File upload section
uploaded_files = st.file_uploader(
    "Upload research papers (PDF)", 
    type=["pdf"], 
    accept_multiple_files=True,
    help="Upload multiple PDF files to analyze them together"
)

if uploaded_files:
    st.info(f"📎 {len(uploaded_files)} files uploaded")
    
    # Processing options
    col1, col2 = st.columns(2)
    with col1:
        processing_option = st.radio(
            "Processing Method",
            ["Quick Analysis", "Deep Analysis"],
            help="Quick: faster but less detailed, Deep: more comprehensive but slower"
        )
    
    with col2:
        summary_style = st.radio(
            "Summary Style",
            ["Academic", "Simplified"],
            help="Academic: maintains technical terms, Simplified: more accessible language"
        )

    if st.button("Analyze Papers", help="Start the analysis process"):
        with st.spinner("Processing documents..."):
            try:
                # Extract text from PDFs
                documents = {}
                key_findings = {}
                
                for file in uploaded_files:
                    text = extract_text_from_pdf(file)
                    documents[file.name] = text
                
                # Process in batches
                all_summaries = {}
                models = get_alternative_models()
                current_model = models["primary"]
                
                st.markdown("### 📊 Processing Documents")
                
                for batch in batch_documents(list(documents.items())):
                    batch_summaries = {}
                    batch_findings = {}
                    
                    progress_bar = st.progress(0)
                    
                    for idx, (filename, content) in enumerate(batch):
                        st.info(f"Processing: {filename}")
                        
                        # Generate regular summary
                        summary = process_document(content, current_model, client)
                        
                        # Extract key findings
                        findings = extract_key_findings(content, current_model, client)
                        
                        if summary and findings:
                            batch_summaries[filename] = summary
                            batch_findings[filename] = findings
                        
                        progress_bar.progress((idx + 1) / len(batch))
                    
                    all_summaries.update(batch_summaries)
                    key_findings.update(batch_findings)
                
                if all_summaries:
                    st.success("Individual document processing complete!")
                    
                    with st.spinner("Creating comprehensive analysis..."):
                        # Generate cross-document summary
                        final_summary = create_cross_document_summary(
                            list(all_summaries.values()),
                            documents,
                            current_model,
                            client
                        )
                        
                        # Validate summaries
                        validation_results = validate_summary(
                            documents,
                            final_summary,
                            all_summaries
                        )
                        
                        # Display validation results
                        if validation_results:
                            st.markdown("### 📊 Summary Validation Metrics")
                            
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.markdown("#### Individual Summaries Scores")
                                scores = validation_results['individual_summaries']['average_scores']
                                st.metric("ROUGE-1", f"{scores['rouge1']:.4f}")
                                st.metric("ROUGE-2", f"{scores['rouge2']:.4f}")
                                st.metric("ROUGE-L", f"{scores['rougeL']:.4f}")
                                st.metric("BLEU", f"{scores['bleu']:.4f}")
                            
                            with col2:
                                st.markdown("#### Cross-Document Summary Scores")
                                scores = validation_results['cross_document_summary']['average_scores']
                                st.metric("ROUGE-1", f"{scores['rouge1']:.4f}")
                                st.metric("ROUGE-2", f"{scores['rouge2']:.4f}")
                                st.metric("ROUGE-L", f"{scores['rougeL']:.4f}")
                                st.metric("BLEU", f"{scores['bleu']:.4f}")
                        
                        # Generate research impact analysis
                        research_impact = analyze_research_impact(
                            list(key_findings.values()),
                            current_model,
                            client
                        )
                        
                        # Display results in organized tabs
                        tab1, tab2, tab3 = st.tabs(["Cross-Document Summary", "Key Research Contributions", "Individual Documents"])
                        
                        with tab1:
                            st.markdown("### 📑 Cross-Document Summary")
                            st.markdown(final_summary)
                        
                        with tab2:
                            st.markdown("### 🔬 Key Research Contributions")
                            st.markdown(research_impact)
                            
                            # Add expandable sections for specific aspects
                            with st.expander("Technical Innovations"):
                                st.markdown("#### Key Technical Contributions")
                                for filename, findings in key_findings.items():
                                    st.markdown(f"**{filename}**")
                                    st.markdown(findings)
                                    st.markdown("---")
                            
                            with st.expander("Methodologies"):
                                st.markdown("#### Research Methodologies")
                                # Extract and display methodology-specific information
                                
                            with st.expander("Future Research Directions"):
                                st.markdown("#### Potential Future Work")
                                # Extract and display future research suggestions
                        
                        with tab3:
                            st.markdown("### 📚 Individual Document Summaries")
                            for filename, summary in all_summaries.items():
                                with st.expander(filename):
                                    st.markdown("#### Summary")
                                    st.markdown(summary)
                                    st.markdown("#### Key Findings")
                                    st.markdown(key_findings[filename])
                        
                        # Save results
                        st.session_state.summaries = all_summaries
                        st.session_state.key_findings = key_findings
                        st.session_state.cross_summary = final_summary
                        st.session_state.research_impact = research_impact
                        
                        # Enhanced download options
                        st.markdown("### 💾 Download Options")
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            if st.button("Download Complete Analysis (PDF)"):
                                complete_analysis = f"""
                                # Research Analysis Report
                                
                                ## Cross-Document Summary
                                {final_summary}
                                
                                ## Key Research Contributions
                                {research_impact}
                                
                                ## Individual Document Analyses
                                {"".join([f'### {filename}\n{summary}\n\nKey Findings:\n{key_findings[filename]}\n\n---\n' 
                                        for filename, summary in all_summaries.items()])}
                                """
                                pdf_buffer = convert_to_pdf(complete_analysis)
                                st.download_button(
                                    "Download Complete PDF",
                                    data=pdf_buffer,
                                    file_name="research_analysis.pdf",
                                    mime="application/pdf"
                                )
                        
                        with col2:
                            if st.button("Download Research Contributions (DOCX)"):
                                docx_buffer = convert_to_docx(research_impact)
                                st.download_button(
                                    "Download Research DOCX",
                                    data=docx_buffer,
                                    file_name="research_contributions.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        
                        with col3:
                            if st.button("Download Summary Only (PDF)"):
                                pdf_buffer = convert_to_pdf(final_summary)
                                st.download_button(
                                    "Download Summary PDF",
                                    data=pdf_buffer,
                                    file_name="research_summary.pdf",
                                    mime="application/pdf"
                                )
            
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

def create_interactive_summary(summary, original_text):
    # Split summary and original text into sentences
    summary_sentences = re.split(r'(?<=[.!?]) +', summary)
    original_sentences = re.split(r'(?<=[.!?]) +', original_text)
    
    # Create a mapping of summary sentences to original text
    mapping = {}
    for i, sum_sent in enumerate(summary_sentences):
        best_match = -1
        highest_ratio = 0
        for j, orig_sent in enumerate(original_sentences):
            ratio = difflib.SequenceMatcher(None, sum_sent, orig_sent).ratio()
            if ratio > highest_ratio:
                highest_ratio = ratio
                best_match = j
        mapping[i] = best_match
    
    # Create HTML for interactive summary
    html = """
    <style>
        .summary-sentence { padding: 2px; border-radius: 3px; cursor: pointer; }
        .summary-sentence:hover { background-color: #e0e0e0; }
        .highlight { background-color: #FFFF00; transition: background-color 0.3s; }
        #summary, #original { height: 300px; overflow-y: auto; padding: 10px; border: 1px solid #ddd; border-radius: 5px; }
        .selected { background-color: #ADD8E6; }
    </style>
    <div id="container" style="display: flex; justify-content: space-between;">
        <div id="summary" style="width: 48%;">
            <h3>Summary</h3>
    """
    for i, sentence in enumerate(summary_sentences):
        html += f'<span class="summary-sentence" id="sum-{i}" onclick="highlight({i})">{sentence} </span>'
    html += '</div>'
    html += '<div id="original" style="width: 48%;">'
    html += '<h3>Original Text</h3>'
    for i, sentence in enumerate(original_sentences):
        html += f'<span id="orig-{i}">{sentence} </span>'
    html += '</div></div>'
    
    # Add JavaScript for interactivity
    html += """
    <script>
        function highlight(i) {
            // Remove previous highlights
            document.querySelectorAll('.highlight, .selected').forEach(el => {
                el.classList.remove('highlight');
                el.classList.remove('selected');
            });
            // Highlight corresponding original sentence
            let origSentence = document.getElementById('orig-' + mapping[i]);
            let sumSentence = document.getElementById('sum-' + i);
            if (origSentence && sumSentence) {
                origSentence.classList.add('highlight');
                sumSentence.classList.add('selected');
                origSentence.scrollIntoView({ behavior: 'smooth', block: 'center' });
            }
        }
        var mapping = """ + json.dumps(mapping) + ";"
    html += '</script>'
    
    return html

def convert_to_pdf(summary):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    c.drawString(72, height - 72, "Summary")
    text_object = c.beginText(72, height - 100)
    for line in summary.split('\n'):
        text_object.textLine(line)
    c.drawText(text_object)
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer

def convert_to_docx(summary):
    doc = Document()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def convert_to_html(summary):
    html = f"""
    <html>
    <head><title>Summary</title></head>
    <body>
    <h1>Summary</h1>
    <p>{summary}</p>
    </body>
    </html>
    """
    return html.encode()

def chunk_text(text, chunk_size=4000, overlap=200):
    """Split text into overlapping chunks."""
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = start + chunk_size
        
        # If this is not the first chunk, include overlap from previous chunk
        if start > 0:
            start = start - overlap
            
        # If this is not the last chunk, try to break at a sentence
        if end < text_length:
            # Try to find the last sentence break in this chunk
            last_period = text[start:end].rfind('.')
            if last_period != -1:
                end = start + last_period + 1
        
        chunks.append(text[start:end])
        start = end
    
    return chunks

# Display summary history
st.markdown("## Summary History")
for entry in st.session_state.summary_history:
    st.markdown(f"**Original Text:** {entry['original']}")
    for model, summary in entry['summaries'].items():
        st.markdown(f"**Summary from {model}:**")
        st.markdown(summary)
    st.markdown("---")

if st.checkbox("Show side-by-side comparison"):
    col1, col2 = st.columns(2)
    models = list(st.session_state.summaries.keys())
    if models:  # Only show comparison if there are summaries
        with col1:
            model1 = st.selectbox("Select first model", models, key="model1_select")
            st.text_area("Summary", st.session_state.summaries[model1], height=300, key="summary1")
        with col2:
            model2 = st.selectbox("Select second model", models, key="model2_select")
            st.text_area("Summary", st.session_state.summaries[model2], height=300, key="summary2")
    else:
        st.info("Generate summaries to enable comparison.")
