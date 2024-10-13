import os
import json
import streamlit as st
from groq import Groq
import PyPDF2
import docx
import io
import streamlit.components.v1 as components
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import base64
import re
import difflib

# Streamlit page configuration
st.set_page_config(
    page_title="LLAMA 3.1. Chat",
    page_icon="🦙",
    layout="centered"
)

working_dir = os.path.dirname(os.path.abspath(__file__))
config_data = json.load(open(f"{working_dir}/config.json"))

# Save the API key to environment variable
os.environ["GROQ_API_KEY"] = config_data["GROQ_API_KEY"]

# Use the API key directly when initializing the client
client = Groq(api_key=os.environ["GROQ_API_KEY"])

# Initialize the summary history if not present
if "summary_history" not in st.session_state:
    st.session_state.summary_history = []

# Initialize summaries if not present
if "summaries" not in st.session_state:
    st.session_state.summaries = {}

# Streamlit page title
st.title("🦙 LLAMA Multi-Model Document Summarizer")

# Function to extract text from PDF
def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Input method selection
input_method = st.radio("Choose input method:", ("Text Input", "File Upload"))

if input_method == "Text Input":
    # Text area for manual input
    user_text = st.text_area("Enter the text you want to summarize:", height=200)
else:
    # File uploader
    uploaded_file = st.file_uploader("Choose a file (PDF or DOCX)", type=["pdf", "docx"])

# Summary length input
summary_length_option = st.radio("Choose summary length option:", ("Word count", "Percentage of original"))

if summary_length_option == "Word count":
    summary_length = st.number_input("Enter desired summary length (in words):", 
                                     min_value=50, max_value=5000, value=200, step=10)
    length_instruction = f"approximately {summary_length} words"
else:
    summary_percentage = st.slider("Select summary length as percentage of original:", 
                                   min_value=10, max_value=50, value=30, step=5)
    length_instruction = f"{summary_percentage}% of the original length"

# Model selection
st.markdown("### Select models to use for summarization:")
use_llama_3_2_11b = st.checkbox("llama-3.2-11b-vision-preview", value=True)
use_llama_3_2_3b = st.checkbox("llama-3.2-3b-preview", value=True)
use_llama_3_1_70b = st.checkbox("llama-3.1-70b-versatile", value=True)
use_llama_3_1_8b = st.checkbox("llama-3.1-8b-instant", value=True)

def create_interactive_summary(summary, original_text):
    # Split summary and original text into sentences
    summary_sentences = re.split(r'(?<=[.!?]) +', summary)
    original_sentences = re.split(r'(?<=[.!?]) +', original_text)
    
    # Create a mapping of summary sentences to original text
    mapping = {}
    for i, sum_sent in enumerate(summary_sentences):
        best_match = -1
        highest_ratio = 0
        for j, orig_sent in enumerate(original_sentences):
            ratio = difflib.SequenceMatcher(None, sum_sent, orig_sent).ratio()
            if ratio > highest_ratio:
                highest_ratio = ratio
                best_match = j
        mapping[i] = best_match
    
    # Create HTML for interactive summary
    html = """
    <style>
        .summary-sentence { padding: 2px; border-radius: 3px; cursor: pointer; }
        .summary-sentence:hover { background-color: #e0e0e0; }
        .highlight { background-color: #FFFF00; transition: background-color 0.3s; }
        #summary, #original { height: 300px; overflow-y: auto; padding: 10px; border: 1px solid #ddd; border-radius: 5px; }
        .selected { background-color: #ADD8E6; }
    </style>
    <div id="container" style="display: flex; justify-content: space-between;">
        <div id="summary" style="width: 48%;">
            <h3>Summary</h3>
    """
    for i, sentence in enumerate(summary_sentences):
        html += f'<span class="summary-sentence" id="sum-{i}" onclick="highlight({i})">{sentence} </span>'
    html += '</div>'
    html += '<div id="original" style="width: 48%;">'
    html += '<h3>Original Text</h3>'
    for i, sentence in enumerate(original_sentences):
        html += f'<span id="orig-{i}">{sentence} </span>'
    html += '</div></div>'
    
    # Add JavaScript for interactivity
    html += """
    <script>
        function highlight(i) {
            // Remove previous highlights
            document.querySelectorAll('.highlight, .selected').forEach(el => {
                el.classList.remove('highlight');
                el.classList.remove('selected');
            });
            // Highlight corresponding original sentence
            let origSentence = document.getElementById('orig-' + mapping[i]);
            let sumSentence = document.getElementById('sum-' + i);
            if (origSentence && sumSentence) {
                origSentence.classList.add('highlight');
                sumSentence.classList.add('selected');
                origSentence.scrollIntoView({ behavior: 'smooth', block: 'center' });
            }
        }
        var mapping = """ + json.dumps(mapping) + ";"
    html += '</script>'
    
    return html

def convert_to_pdf(summary):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    c.drawString(72, height - 72, "Summary")
    text_object = c.beginText(72, height - 100)
    for line in summary.split('\n'):
        text_object.textLine(line)
    c.drawText(text_object)
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer

def convert_to_docx(summary):
    doc = Document()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def convert_to_html(summary):
    html = f"""
    <html>
    <head><title>Summary</title></head>
    <body>
    <h1>Summary</h1>
    <p>{summary}</p>
    </body>
    </html>
    """
    return html.encode()

if st.button("Summarize"):
    if input_method == "Text Input" and user_text:
        text_to_summarize = user_text
    elif input_method == "File Upload" and uploaded_file is not None:
        # Process uploaded file
        if uploaded_file.type == "application/pdf":
            text_to_summarize = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text_to_summarize = extract_text_from_docx(io.BytesIO(uploaded_file.getvalue()))
    else:
        st.warning("Please enter text or upload a file to summarize.")
        st.stop()

    # List of selected models
    selected_models = []
    if use_llama_3_2_11b:
        selected_models.append("llama-3.2-11b-vision-preview")
    if use_llama_3_2_3b:
        selected_models.append("llama-3.2-3b-preview")
    if use_llama_3_1_70b:
        selected_models.append("llama-3.1-70b-versatile")
    if use_llama_3_1_8b:
        selected_models.append("llama-3.1-8b-instant")

    if not selected_models:
        st.warning("Please select at least one model for summarization.")
        st.stop()

    st.session_state.summaries = {}

    for model in selected_models:
        messages = [
            {"role": "system", "content": """You are an expert summarizer with a keen ability to distill complex information into clear, concise summaries. Your task is to create a comprehensive yet concise summary of the given text. Please follow these guidelines:

1. Identify and include the main ideas and key points of the original text.
2. Maintain the original tone and style of the text where appropriate.
3. Correct any obvious OCR errors in your summary, but preserve the meaning of the original text.
4. Ensure the summary is coherent and flows logically.
5. Use clear and concise language, avoiding unnecessary jargon or complexity.
6. If the text contains multiple sections or topics, provide a balanced representation of each.
7. Include relevant examples or data points if they are crucial to understanding the main ideas.
8. Avoid including your own opinions or interpretations; stick to the information provided in the original text.
9. If the text contains any actionable items or conclusions, be sure to include them in the summary.

Your goal is to create a summary that allows the reader to quickly grasp the essence of the original text without losing important details."""},
            {"role": "user", "content": f"Please summarize the following text in {length_instruction}:\n\n{text_to_summarize[:4000]}"}
        ]

        response = client.chat.completions.create(
            model=model,
            messages=messages
        )

        st.session_state.summaries[model] = response.choices[0].message.content

    # Add the summaries to the history
    st.session_state.summary_history.append({
        "original": text_to_summarize[:5000] + "..." if len(text_to_summarize) > 5000 else text_to_summarize,
        "summaries": st.session_state.summaries
    })

    # Display the summaries
    for model, summary in st.session_state.summaries.items():
        st.markdown(f"**Summary from {model}:**")
        st.markdown(summary)
        st.markdown("---")
        
        # Create and display interactive summary
        interactive_summary = create_interactive_summary(summary, text_to_summarize)
        components.html(interactive_summary, height=600)
        
        # Add download options
        output_format = st.selectbox(f"Select output format for {model}:", ["PDF", "DOCX", "HTML"])
        
        if st.button(f"Download Summary for {model}"):
            if output_format == "PDF":
                buf = convert_to_pdf(summary)
                st.download_button(
                    label="Download PDF",
                    data=buf,
                    file_name=f"summary_{model}.pdf",
                    mime="application/pdf"
                )
            elif output_format == "DOCX":
                buf = convert_to_docx(summary)
                st.download_button(
                    label="Download DOCX",
                    data=buf,
                    file_name=f"summary_{model}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:  # HTML
                buf = convert_to_html(summary)
                b64 = base64.b64encode(buf).decode()
                href = f'<a href="data:text/html;base64,{b64}" download="summary_{model}.html">Download HTML</a>'
                st.markdown(href, unsafe_allow_html=True)

# Display summary history
st.markdown("## Summary History")
for entry in st.session_state.summary_history:
    st.markdown(f"**Original Text:** {entry['original']}")
    for model, summary in entry['summaries'].items():
        st.markdown(f"**Summary from {model}:**")
        st.markdown(summary)
    st.markdown("---")

if st.checkbox("Show side-by-side comparison"):
    col1, col2 = st.columns(2)
    models = list(st.session_state.summaries.keys())
    if models:  # Only show comparison if there are summaries
        with col1:
            model1 = st.selectbox("Select first model", models, key="model1_select")
            st.text_area("Summary", st.session_state.summaries[model1], height=300, key="summary1")
        with col2:
            model2 = st.selectbox("Select second model", models, key="model2_select")
            st.text_area("Summary", st.session_state.summaries[model2], height=300, key="summary2")
    else:
        st.info("Generate summaries to enable comparison.")
